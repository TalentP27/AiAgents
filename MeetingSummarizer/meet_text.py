import assemblyai as aai
from dotenv import load_dotenv
import os
from agno.agent import Agent
from agno.models.google import Gemini
from agno.models.openai import OpenAIChat
from dotenv import load_dotenv
import os
from agno.tools.file import FileTools
from docx import Document
from agno.tools.gmail import GmailTools

#loading the env
load_dotenv()
id=os.getenv("id")
api_key=os.getenv("api_key_gemini_v2")
id_openai=os.getenv("id_openai")
api_key_openai = os.getenv("api_key_openai")

aai.settings.api_key = os.getenv("ASSEMBLYAI_API_KEY")


audio_file =r"D:\Agno\Meeting_text\video\videoplayback.m4a"
# assembly ai configuration
config = aai.TranscriptionConfig(
  speaker_labels=True,
)
# transcribe the audio
transcript = aai.Transcriber().transcribe(audio_file, config)

# store it in a txt file
with open("Conv_to_text.txt", "w", encoding="utf-8") as file:
    for utterance in transcript.utterances:
        file.write(f"Speaker {utterance.speaker}: {utterance.text}\n")
        # print(f"Speaker {utterance.speaker}: {utterance.text}")
       

with open("Conv_to_text.txt","r") as file:
    full_conv=file.read()

# agent for summarizing the full conversation
summarizer = Agent(
  name="summarizer agent",
  model=Gemini(id=id,api_key=api_key),
#   model=OpenAIChat(id=id_openai, api_key=api_key_openai),
  description="""
    An intelligent summarization agent that reads through a multi-speaker conversation transcript and extracts meaningful insights. It is designed to provide a concise yet informative summary, highlight key discussion points, and identify specific examples, decisions, or action items mentioned.
    """,

  instructions="""
    You are a summarization expert. Your task is to analyze a detailed transcript of a conversation between multiple speakers. Follow these steps:

    1. Read the full transcript carefully.
    2. Generate a clear and concise summary of the overall conversation.
    3. Identify and list the main topics or sections discussed.
    4. Extract any examples used, decisions made, or action items assigned.
    5. Format your output in a structured and easy-to-read way.

    Make sure to remain neutral and accurate in tone. Preserve speaker intent without adding any extra interpretation.
    """
    ,
  expected_output="""
    Your response should follow this structure:

    ### Summary
    A 3–6 sentence paragraph summarizing the entire conversation.

    ### Key Topics Discussed
    - Topic 1
    - Topic 2
    - Topic 3
    ...

    ### Notable Examples
    - "Example or story quoted from speaker"
    - ...

    ### Decisions or Action Items
    -  Speaker A will follow up on...
    -  Decision to proceed with...
    """
    ,
    tools=[FileTools()],
    markdown=True,

)


# summarizer.print_response("summarize the conversation for me on the file Conv_to_text.txt",stream=True, markdown=True)

summary = summarizer.run("summarize the conversation for me on the file Conv_to_text.txt")

# Extract the text content from the RunResponse object
summary_text = summary.content if hasattr(summary, "content") else str(summary)

# reformating the output for a docx file
doc = Document()
doc.add_heading("Meeting Summary", 0)

for line in summary_text.split('\n'):
    line = line.strip()
    if not line:
        continue
    if line.startswith("### "): #title
        doc.add_heading(line[4:], level=2)
    elif line.startswith("# "):
        doc.add_heading(line[2:], level=1)
    elif line.startswith("- "):
        doc.add_paragraph(line[2:], style='List Bullet')
    else:
        doc.add_paragraph(line)

doc.save("Meeting_Summary.docx")
print("Summary saved to Meeting_Summary.docx")

# Agent to send the docx file through gmail
Gmail_sender = Agent(
    name="Gmail mail sender",
    model=Gemini(id=id,api_key=api_key),
    # model=OpenAIChat(id=id_openai, api_key=api_key_openai),
    tools=[GmailTools(credentials_path=r"D:\Agno\Meeting_text\client_secret_14840150298-jkg1i33kl9vgf0j1d6ro1aqu0va9e3qg.apps.googleusercontent.com.json"), FileTools()],
    instructions="""
    Your role is to send the 'Meeting_Summary.docx' document to the list of recipients provided in the user's message.
    
    1. Craft a clear, professional email message based on the meeting summary content.
    2. Use the file 'Meeting_Summary.docx' as an attachment.
    3. Accept a list of one or more email addresses to send the message to.
    4. Include a meaningful subject (e.g., "Meeting Summary – [Insert Date or Topic]").
    5. Use GmailTools to send the email with the attachment.
    6. Confirm that the email was successfully sent or report any errors.
    
    If the user provides additional context (e.g., meeting title, purpose, or message body preferences), incorporate that into the email.
       """,
       show_tool_calls=True,
)


response = Gmail_sender.run(
    "Send the Meeting_Summary.docx to younessichen11@gmail.com"
)
print(response)